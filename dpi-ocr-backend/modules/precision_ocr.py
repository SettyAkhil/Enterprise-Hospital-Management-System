# modules/precision_ocr.py  — Production OCR Engine v2.0
"""
Production-level OCR pipeline using gemma4:26b vision model.
Features:
  - 5 adaptive image preprocessing strategies with auto-retry
  - Multi-page PDF support (processes all pages, merges results)
  - Image upscaling for low-resolution scans
  - Robust response parsing with fallback chain
  - Per-step error surfacing (no silent failures)
  - Archive vault integration
  - Download in MD / PDF / DOCX formats
"""
import time
import io
import asyncio
import json
import logging
import fitz          # PyMuPDF
import re
import math
import numpy as np
import cv2
from PIL import Image, ImageEnhance, ImageOps, ImageFilter
from openai import OpenAI
import base64
import markdown
from xhtml2pdf import pisa
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import time
import pandas as pd
from core.config import settings
from modules.unified_resolver import resolve_entities_in_text
from modules.layout_detector import DocLayoutDetector
from modules.reading_order import ReadingOrderEngine
from modules.region_ocr import AsyncRegionOCR
from modules.table_extractor import TableExtractor
from modules.medical_corrector import MedicalCorrector
from modules.confidence_engine import ConfidenceEngine

# Bounded timeout (SDK default is 600s) — a hung/overloaded vLLM server should
# fail one call fast and let Celery's retry policy handle it, not silently
# block a worker for up to 10 minutes per call.
# Always settings.VLLM_BASE_URL, never a hardcoded "localhost:8700" — a
# hardcoded host string here silently breaks in any containerized deployment.
client = OpenAI(base_url=settings.VLLM_BASE_URL, api_key="EMPTY", timeout=90.0)
logger = logging.getLogger(__name__)

# ---------------------------------------------------------------------------
# 0.  PAGE ORIENTATION DETECTION (0/90/180/270)
# ---------------------------------------------------------------------------
def detect_page_orientation(img: Image.Image) -> int:
    """
    Detects page orientation. First attempts pytesseract OSD: if conf >= 2.0, uses OSD angle (0/90/180/270).
    If OSD is unavailable or conf < 2.0, falls back to row-projection variance:
      - portrait axis -> 0 vs 180 vLLM confidence probe
      - quarter-turn axis -> returns 0 and sets img.info["needs_manual_orientation"] = True.
    """
    try:
        import pytesseract
        osd = pytesseract.image_to_osd(img, output_type=pytesseract.Output.DICT)
        angle = int(osd.get("rotate", 0))
        conf = float(osd.get("orientation_conf", 0.0))
        logger.info(f"Pytesseract OSD orientation detection: angle={angle}, conf={conf:.2f}")
        if conf >= 2.0:
            return angle
    except Exception as e:
        logger.warning(f"Pytesseract OSD detection unavailable or failed: {e}")


    arr = np.array(img.convert("L"))


    def _row_projection_variance(rotation_deg: int) -> float:
        rotated = np.rot90(arr, k=rotation_deg // 90)
        return float(np.var(rotated.sum(axis=1)))

    portrait_axis = _row_projection_variance(0) >= _row_projection_variance(90)
    if not portrait_axis:
        logger.info("Page appears rotated a quarter-turn (90/270) — flagging manual orientation warning.")
        if hasattr(img, "info") and isinstance(img.info, dict):
            img.info["needs_manual_orientation"] = True
        return 0

    candidates = (0, 180)

    w, h = img.size
    crop_w, crop_h = min(w, 700), min(h, 300)
    cx, cy = w // 2, h // 2
    center_crop = img.crop(
        (cx - crop_w // 2, cy - crop_h // 2, cx + crop_w // 2, cy + crop_h // 2)
    )

    scores = {}
    for rot in candidates:
        test_img = center_crop.rotate(-rot, expand=True) if rot else center_crop
        scores[rot] = _vllm_orientation_confidence(test_img)

    ranked = sorted(candidates, key=lambda r: scores[r], reverse=True)
    winner, runner_up = ranked[0], ranked[1]
    MIN_ABSOLUTE_CONFIDENCE = 0.5
    MIN_MARGIN_RATIO = 1.3
    if scores[winner] < MIN_ABSOLUTE_CONFIDENCE or scores[winner] < MIN_MARGIN_RATIO * max(scores[runner_up], 1e-6):
        return 0
    return winner



def _vllm_orientation_confidence(img: Image.Image) -> float:
    """Mean per-token log-prob (exponentiated) of a short vLLM read of `img`,
    used only to compare orientation candidates against each other.

    Deliberately creates its own client per call rather than using the
    module-level `client` above: that client is constructed at import time,
    and Celery's prefork worker model forks child processes *after* the
    module (and its at-import-time HTTP client/connection pool) is loaded —
    reusing a pre-fork httpx client from a forked child is a known way to
    deadlock silently. Confirmed by reproduction: page tasks hung for
    ~10 minutes (the OpenAI SDK's default request timeout) with zero
    progress the first time this code path ran under the real Celery worker,
    vanishing once each call got its own client instead of sharing the
    pre-fork one.
    """
    try:
        buf = io.BytesIO()
        img.convert("RGB").save(buf, format="JPEG", quality=85)
        b64 = base64.b64encode(buf.getvalue()).decode("utf-8")
        local_client = OpenAI(base_url=settings.VLLM_BASE_URL, api_key="EMPTY", timeout=30.0)
        resp = local_client.chat.completions.create(
            model="qwen2.5-vl-7b",
            messages=[{
                "role": "user",
                "content": [
                    {"type": "text", "text": "Read the text in this image exactly as written."},
                    {"type": "image_url", "image_url": {"url": f"data:image/jpeg;base64,{b64}"}},
                ],
            }],
            temperature=0,
            max_tokens=40,
            logprobs=True,
            top_logprobs=1,
        )
        content = resp.choices[0].logprobs.content if resp.choices[0].logprobs else None
        if not content:
            return 0.0
        return float(np.exp(np.mean([t.logprob for t in content])))
    except Exception as e:
        logger.warning(f"Orientation confidence probe failed: {e}")
        return 0.0


def apply_orientation(img: Image.Image) -> Image.Image:
    """Detects and corrects full-rotation page orientation. No-op (cheap) for
    the common already-upright case; returns the input unchanged then."""
    rotation = detect_page_orientation(img)
    if rotation == 0:
        return img
    logger.info(f"Correcting page rotation by {rotation} degrees")
    return img.rotate(-rotation, expand=True)


# ---------------------------------------------------------------------------
# 0b. BLANK PAGE DETECTION
# ---------------------------------------------------------------------------
def is_blank_page(img: Image.Image, ink_ratio_threshold: float = 0.01) -> bool:
    """
    Cheap (no model call) check for near-empty separator/blank pages, common
    in real hospital scan batches. Thresholds against the page's own median
    background level (works for off-white/gray scans, not just pure white)
    and measures what fraction of pixels are meaningfully darker than that.

    Tuned against real documents in uploads/: genuine content pages measured
    6-17% ink coverage; a synthetic blank page (with scanner noise) measured
    0%, and a near-blank page with only a small stamp/mark measured 0.4% —
    both comfortably under this threshold, with a wide margin to real content.
    """
    gray = np.array(img.convert("L"), dtype=np.float32)
    bg_level = np.median(gray)
    ink_mask = gray < (bg_level - 30)
    ink_ratio = float(ink_mask.mean())
    return ink_ratio < ink_ratio_threshold

# ---------------------------------------------------------------------------
# 1.  IMAGE PREPROCESSING STRATEGIES
MIN_DIM = 600   # minimum pixel dimension; smaller images are upscaled

def _upscale_if_small(img: Image.Image) -> Image.Image:
    """Upscale image if either dimension is below MIN_DIM."""
    w, h = img.size
    if min(w, h) < MIN_DIM:
        scale = MIN_DIM / min(w, h)
        # Cap scale factor to prevent line crops (e.g. 800x40) from becoming massively oversized
        if scale > 2.5:
            scale = 2.5
        # Ensure the maximum dimension doesn't exceed 3000 after upscaling
        if max(w, h) * scale > 3000:
            scale = 3000 / max(w, h)
        
        if scale > 1.0:
            new_w, new_h = int(w * scale), int(h * scale)
            img = img.resize((new_w, new_h), Image.BILINEAR)
    return img

def strategy_original(img: Image.Image) -> Image.Image:
    """Strategy 0 — Send the image with minimal changes (fix orientation only)."""
    img = ImageOps.exif_transpose(img)
    img = _upscale_if_small(img)
    return img.convert("RGB")

def strategy_mild_enhance(img: Image.Image) -> Image.Image:
    """Strategy 1 — Mild sharpening + contrast boost. Best for printed forms."""
    img = ImageOps.exif_transpose(img)
    img = _upscale_if_small(img)
    img = img.convert("RGB")
    img = ImageEnhance.Sharpness(img).enhance(1.8) 
    img = ImageEnhance.Contrast(img).enhance(1.5)
    return img

def strategy_grayscale_boost(img: Image.Image) -> Image.Image:
    """Strategy 2 — Grayscale + moderate contrast. Good for photocopies."""
    img = ImageOps.exif_transpose(img)
    img = _upscale_if_small(img)
    gray = img.convert("L")
    gray = ImageEnhance.Contrast(gray).enhance(1.8)
    gray = ImageEnhance.Sharpness(gray).enhance(2.0)
    return gray.convert("RGB")

def strategy_adaptive_threshold(img: Image.Image) -> Image.Image:
    """Strategy 3 — Near-binarization for faded handwriting on white background."""
    img = ImageOps.exif_transpose(img)
    img = _upscale_if_small(img)
    gray = img.convert("L")
    # Auto-level: stretch histogram
    arr = np.array(gray, dtype=np.float32)
    p2, p98 = np.percentile(arr, 2), np.percentile(arr, 98)
    if p98 > p2:
        arr = np.clip((arr - p2) / (p98 - p2) * 255, 0, 255)
    gray = Image.fromarray(arr.astype(np.uint8))
    gray = ImageEnhance.Contrast(gray).enhance(2.2)
    # Unsharp mask to make thin strokes pop
    gray = gray.filter(ImageFilter.UnsharpMask(radius=1, percent=150, threshold=3))
    return gray.convert("RGB")

def strategy_denoised(img: Image.Image) -> Image.Image:
    """Strategy 4 — Denoise first then enhance. For camera photos of documents."""
    img = ImageOps.exif_transpose(img)
    img = _upscale_if_small(img)
    img = img.convert("RGB")
    # Median filter removes camera noise
    img = img.filter(ImageFilter.MedianFilter(size=3))
    img = ImageEnhance.Contrast(img).enhance(1.6)
    img = ImageEnhance.Sharpness(img).enhance(2.0)
    return img

def _deskew_cv(img_bgr: np.ndarray) -> np.ndarray:
    """Hough-line-based skew detection + rotation correction."""
    gray = cv2.cvtColor(img_bgr, cv2.COLOR_BGR2GRAY)
    edges = cv2.Canny(gray, 50, 150, apertureSize=3)
    lines = cv2.HoughLinesP(edges, 1, np.pi / 180, 100, minLineLength=100, maxLineGap=10)
    if lines is None:
        return img_bgr
    # reshape(-1, 4), not lines[:, 0]: HoughLinesP's output shape is (N, 1, 4)
    # on some opencv-python-headless versions and (N, 4) on others (hit this
    # for real — a fresh `pip install` picked up a newer version than the
    # long-running host venv had, and lines[:, 0] silently returned scalars
    # instead of (x1,y1,x2,y2) tuples on the (N, 4) shape, breaking deskew
    # for every image). reshape(-1, 4) is correct for both.
    angles = [np.degrees(np.arctan2(y2 - y1, x2 - x1)) for (x1, y1, x2, y2) in lines.reshape(-1, 4)]
    if not angles:
        return img_bgr
    median_angle = np.median(angles)
    # Only correct noticeable skew; skip near-0 (already straight) and near-90 (rotated page, different problem).
    if not (0.5 < abs(median_angle) < 15):
        return img_bgr
    h, w = img_bgr.shape[:2]
    m = cv2.getRotationMatrix2D((w // 2, h // 2), median_angle, 1.0)
    return cv2.warpAffine(img_bgr, m, (w, h), flags=cv2.INTER_CUBIC, borderMode=cv2.BORDER_REPLICATE)


def strategy_deskew_clahe_denoise(img: Image.Image) -> Image.Image:
    """Strategy 5 — Deskew + CLAHE contrast + non-local-means denoise. For crooked
    camera photos / fax-quality scans where the other strategies leave visible tilt
    or uneven lighting."""
    img = ImageOps.exif_transpose(img)
    img = _upscale_if_small(img)
    bgr = cv2.cvtColor(np.array(img.convert("RGB")), cv2.COLOR_RGB2BGR)
    bgr = _deskew_cv(bgr)

    lab = cv2.cvtColor(bgr, cv2.COLOR_BGR2LAB)
    l, a, b = cv2.split(lab)
    clahe = cv2.createCLAHE(clipLimit=2.0, tileGridSize=(8, 8))
    l = clahe.apply(l)
    bgr = cv2.cvtColor(cv2.merge((l, a, b)), cv2.COLOR_LAB2BGR)
    bgr = cv2.fastNlMeansDenoisingColored(bgr, None, 10, 10, 7, 21)

    return Image.fromarray(cv2.cvtColor(bgr, cv2.COLOR_BGR2RGB))


STRATEGIES = [
    ("Original",              strategy_original),
    ("Mild Enhancement",      strategy_mild_enhance),
    ("Grayscale Boost",       strategy_grayscale_boost),
    ("Adaptive Threshold",    strategy_adaptive_threshold),
    ("Denoise + Enhance",     strategy_denoised),
    ("Deskew + CLAHE",        strategy_deskew_clahe_denoise),
]

def assess_input_quality(img: Image.Image) -> dict:
    """
    Assesses image resolution, estimated DPI, stroke_px, and band_px.
    Logs q["class"], stroke_px, and band_px.
    """
    w, h = img.size
    est_dpi = int((min(w, h) / 8.5))
    stroke_px = max(1.0, min(w, h) / 1000.0 * 2.0)
    band_px = max(10.0, h / 40.0)

    quality_class = "LEGIBLE" if est_dpi >= 200 and stroke_px >= 2.0 else "DEGRADED"

    logger.info(
        f"Input Quality Assessment: class={quality_class}, est_dpi={est_dpi}, "
        f"stroke_px={stroke_px:.2f}, band_px={band_px:.2f}, dimensions={w}x{h}"
    )

    return {
        "class": quality_class,
        "est_dpi": est_dpi,
        "stroke_px": stroke_px,
        "band_px": band_px,
        "width": w,
        "height": h
    }


def img_to_bytes(img: Image.Image, quality: int = 92) -> bytes:
    buf = io.BytesIO()
    img.save(buf, format="JPEG", quality=quality)
    return buf.getvalue()

def enforce_template_labels(md: str, template) -> tuple[str, list[str]]:
    """
    Enforces ground-truth template labels on an output markdown string.
    Removes any invented/hallucinated rows and inserts missing template rows.
    Returns (repaired_md, violations).
    """
    if not md or not template or not getattr(template, 'rows', None):
        return md, []

    lines = md.splitlines()
    table_start_idx = None
    table_lines = []

    for idx, line in enumerate(lines):
        if line.strip().startswith("|") and "|" in line:
            if table_start_idx is None:
                table_start_idx = idx
            table_lines.append(line)

    if not table_lines or len(table_lines) < 2:
        return md, []

    header = table_lines[0]
    delimiter = table_lines[1]
    header_cells = [c.strip() for c in header.strip("|").split("|")]

    check_item_idx = 2 if len(header_cells) > 2 and "check item" in header_cells[2].lower() else 0

    actual_rows_by_label = {}
    violations = []

    for line in table_lines[2:]:
        cells = [c.strip() for c in line.strip("|").split("|")]
        if len(cells) > check_item_idx:
            label = cells[check_item_idx]
            if label in template.labels:
                actual_rows_by_label[label] = cells
            else:
                violations.append(label)

    n_entry_cols = len(header_cells) - 3 if len(header_cells) > 3 else max(1, len(header_cells) - 1)
    repaired_table_lines = [header, delimiter]

    for tmpl_row in template.rows:
        lbl = tmpl_row.label
        if lbl in actual_rows_by_label:
            row_cells = actual_rows_by_label[lbl]
        else:
            row_cells = [tmpl_row.section, tmpl_row.group, lbl] + [""] * n_entry_cols

        repaired_table_lines.append("| " + " | ".join(row_cells) + " |")

    non_table_before = lines[:table_start_idx]
    non_table_after = lines[table_start_idx + len(table_lines):]

    repaired_md = "\n".join(non_table_before + repaired_table_lines + non_table_after)
    return repaired_md, violations


# ---------------------------------------------------------------------------
# 2.  RESPONSE CLEANING
# ---------------------------------------------------------------------------
def clean_output(raw: str) -> str:

    """
    Remove model control tokens and code fences.
    Falls back to the raw text if cleaning removes everything.
    Preserves medical symbols like < > in reference ranges.
    """
    if not raw:
        return ""

    text = raw

    # Strip content before <|text|> markers (some model variants)
    if "<|text|>" in text:
        text = text.split("<|text|>")[-1]
    if "</|text|>" in text:
        text = text.split("</|text|>")[0]

    # Remove thinking blocks — both formats used by different model variants
    text = re.sub(r'<\|think\|>.*?</\|think\|>', '', text, flags=re.DOTALL)   # pipe format
    text = re.sub(r'<think>.*?</think>',           '', text, flags=re.DOTALL)   # plain format

    # Remove code fences
    text = re.sub(r'```[a-z]*\n?', '', text)
    text = text.replace('```', '')

    epilogue_pattern = re.compile(
        r'(?:\n+|(?<=\|)\s*\n)'
        r'(?:This|The above|Note:|Please note|I have|I\'ve|All text|The transcription|'
        r'The extracted|Here is|Above is)\b'
        r'[^\n]{0,400}?'
        r'(?:transcription|extracted|rules|accurately|as instructed|'
        r'per the instructions|marked|fidelity|provided|captured?)\b[^\n]*$',
        re.IGNORECASE
    )

    for _ in range(3):
        new_text = epilogue_pattern.sub('', text)
        if new_text == text:
            break
        text = new_text

    cleaned = text.strip()

    # Safety fallback: if stripping removed EVERYTHING, return raw minus fences
    if not cleaned:
        fallback = re.sub(r'```[a-z]*\n?', '', raw).replace('```', '').strip()
        return fallback

    return cleaned


# ---------------------------------------------------------------------------
# 3.  PDF MULTI-PAGE LOADER
# ---------------------------------------------------------------------------
def load_pdf_page_count(file_bytes: bytes) -> int:
    """Page count without rendering — cheap, used to size the Celery chord fan-out."""
    doc = fitz.open(stream=file_bytes, filetype="pdf")
    count = len(doc)
    doc.close()
    return count


def load_pdf_page(file_bytes: bytes, page_idx: int) -> Image.Image:
    """Render a single PDF page as a PIL image at 300 DPI. Used by the per-page
    Celery task so a worker only ever holds one rendered page in memory at a
    time, regardless of document length (contrast with load_pdf_pages below,
    which is a convenience for small/synchronous callers only)."""
    doc = fitz.open(stream=file_bytes, filetype="pdf")
    page = doc.load_page(page_idx)
    pix = page.get_pixmap(dpi=300)
    img = Image.open(io.BytesIO(pix.tobytes("png")))
    img.load()  # decode before closing the underlying doc/buffer
    doc.close()
    return img


def load_pdf_page_thumbnail(file_bytes: bytes, page_idx: int, dpi: int = 72) -> Image.Image:
    """Cheap low-DPI render of a single page, used only for perceptual-hash
    duplicate detection at dispatch time (see workers/celery_app.py) — avoids
    paying full 300 DPI render cost twice (once to hash, once to OCR) for
    every page just to detect the minority that are duplicates."""
    doc = fitz.open(stream=file_bytes, filetype="pdf")
    page = doc.load_page(page_idx)
    pix = page.get_pixmap(dpi=dpi)
    img = Image.open(io.BytesIO(pix.tobytes("png")))
    img.load()
    doc.close()
    return img


def page_phash(img: Image.Image) -> str:
    """Perceptual hash for near-duplicate page detection (tolerant to minor
    re-scan/re-compress differences that an exact byte/MD5 hash would miss)."""
    import imagehash
    return str(imagehash.phash(img))


def load_pdf_pages(file_bytes: bytes, max_pages: int | None = None) -> list[Image.Image]:
    """Extract pages from a PDF as PIL images at 300 DPI (Balanced for Qwen).

    Loads every requested page into memory at once — fine for the synchronous/
    single-process callers (tests, small documents) but NOT what the Celery OCR
    pipeline uses for large PDFs; that path calls load_pdf_page_count/
    load_pdf_page one page at a time instead. No page cap by default."""
    doc = fitz.open(stream=file_bytes, filetype="pdf")
    total = len(doc) if max_pages is None else min(len(doc), max_pages)
    pages = []
    for i in range(total):
        page = doc.load_page(i)
        pix  = page.get_pixmap(dpi=300)
        pages.append(Image.open(io.BytesIO(pix.tobytes("png"))))
    doc.close()
    return pages

# ---------------------------------------------------------------------------
# 5.  EXPORT GENERATORS  (PDF / DOCX)
# ---------------------------------------------------------------------------
def _is_table_separator(line: str) -> bool:
    """
    True for a Markdown table separator/alignment row (e.g. '|---|---|',
    '|:--:|-|'). Shared by every download format (PDF/DOCX/Excel/JSON) so a
    row is classified the same way everywhere — before this was unified,
    DOCX used a bare `'---' in line` check (misses a minimal '|-|-|-|'
    separator, and can wrongly swallow a real data row that merely contains
    the substring '---') while Excel/JSON didn't allow ':' (misreads a
    colon-alignment separator like '|:---:|' as a real data row).
    """
    return bool(re.fullmatch(r'[\s|:\-]+', line)) and '-' in line


def prune_empty_table_rows(text: str) -> str:
    """
    Walks pipe-led lines and removes any row whose cells are ALL empty,
    EXCEPT never remove a separator row and never remove the header row.
    Return the text unchanged if it contains no tables.
    """
    lines = text.split('\n')
    if not any(line.strip().startswith('|') for line in lines):
        return text

    out = []
    i, n = 0, len(lines)
    while i < n:
        if lines[i].strip().startswith('|'):
            j = i
            block = []
            while j < n and lines[j].strip().startswith('|'):
                block.append(lines[j])
                j += 1

            pruned = []
            header_seen = False
            for line in block:
                raw = line.strip()
                if _is_table_separator(raw):
                    pruned.append(line)
                    continue

                cells = [c.strip() for c in raw.strip('|').split('|')]
                if not header_seen:
                    pruned.append(line)
                    header_seen = True
                else:
                    if any(c for c in cells):
                        pruned.append(line)

            out.extend(pruned)
            i = j
        else:
            out.append(lines[i])
            i += 1
    return '\n'.join(out)


def deduplicate_markdown_tables(text: str) -> str:
    """
    Deduplicates markdown tables that overlap >= 80%.
    """
    lines = text.split('\n')
    if not any(line.strip().startswith('|') for line in lines):
        return text

    out = []
    i, n = 0, len(lines)
    seen_tables = []
    while i < n:
        if lines[i].strip().startswith('|'):
            j = i
            table_lines = []
            while j < n and lines[j].strip().startswith('|'):
                table_lines.append(lines[j])
                j += 1

            is_dup = False
            lines1 = set(table_lines)
            for prev in seen_tables:
                lines2 = set(prev)
                if lines1 and lines2:
                    overlap = len(lines1.intersection(lines2)) / max(len(lines1), len(lines2))
                    if overlap >= 0.8:
                        is_dup = True
                        break
            if not is_dup:
                seen_tables.append(table_lines)
                out.extend(table_lines)
            i = j
        else:
            out.append(lines[i])
            i += 1
    return '\n'.join(out)


def _normalize_markdown_tables(text: str) -> str:
    """
    Rebuilds every block of consecutive '|'-led lines into a strict, valid
    Markdown table (matching column counts, a proper '|---|' separator row)
    before handing text to python-markdown's 'tables' extension.
    """
    lines = text.split('\n')
    out = []
    i, n = 0, len(lines)
    while i < n:
        if lines[i].strip().startswith('|'):
            j = i
            block = []
            while j < n and lines[j].strip().startswith('|'):
                block.append(lines[j].strip())
                j += 1

            rows = []
            for raw in block:
                if _is_table_separator(raw):
                    continue  # drop the model's own (possibly malformed) separator row
                cells = [c.strip() for c in raw.strip('|').split('|')]
                # Skip rows where every cell is empty before appending to rows (except header)
                if rows and not any(c for c in cells):
                    continue
                rows.append(cells)


            if rows:
                col_count = len(rows[0])
                fixed = []
                for r in rows:
                    if len(r) < col_count:
                        r = r + [''] * (col_count - len(r))
                    elif len(r) > col_count:
                        r = r[:col_count]
                    # xhtml2pdf collapses a column's width to ~0 (ignoring any
                    # explicit/fixed width) if any cell in it is truly empty,
                    # so a blank field — which the LDSL/Healmax blueprints
                    # explicitly produce for missing data — must never reach
                    # the PDF renderer as an empty cell.
                    r = [c if c.strip() else '&nbsp;' for c in r]
                    fixed.append(r)
                out.append('| ' + ' | '.join(fixed[0]) + ' |')
                out.append('|' + '|'.join(['---'] * col_count) + '|')
                out.extend('| ' + ' | '.join(r) + ' |' for r in fixed[1:])
            i = j
        else:
            out.append(lines[i])
            i += 1
    return '\n'.join(out)


def _widest_table_column_count(text: str) -> int:
    """
    Max column count across every Markdown table row in the text (post
    _normalize_markdown_tables, so rows are already well-formed). Landscape
    orientation should key off this, not merely "does a table exist" — a
    narrow 4-column LDSL table fits comfortably in portrait, and forcing it
    into landscape actively hurts: A4 landscape has *less* vertical room
    (595pt) than portrait (842pt), so the same row count that fit on one
    portrait page spills onto a second landscape page for no benefit. A wide
    9-column Healmax table genuinely needs the extra horizontal room though.
    """
    max_cols = 0
    for line in text.split('\n'):
        line = line.strip()
        if line.startswith('|') and not _is_table_separator(line):
            max_cols = max(max_cols, len(line.strip('|').split('|')))
    return max_cols


_META_FIELD_RE = re.compile(r'^\*{0,2}([^*:|]+):\*{0,2}\s*(.*)$')


def _is_meta_field_segment(segment: str) -> bool:
    return bool(_META_FIELD_RE.match(segment.strip()))


def _render_meta_group(lines: list[str]) -> str:
    """
    Renders one or more consecutive 'Label: Value | Label: Value' lines
    (patient info, doctor/date fields, footer fields — bold or not) as a
    single borderless table with one row per line, instead of plain run-on
    text with literal '|' characters. A real HTML table is used (rather than
    e.g. flexbox) because xhtml2pdf only supports a small, table/block-based
    subset of CSS2.1 — flex layout is silently dropped.

    The divider rule goes on the LAST row's cells only, never on the <table>
    or a wrapping <div>: xhtml2pdf repeats a border set on a multi-child
    block container onto every child row instead of drawing it once at the
    bottom (verified empirically), so putting it on the container would draw
    a line under every single field row instead of one line under the group.
    (A leading rule was tried for footer groups too, to box them like a
    signature bar, but immediately following another group's own trailing
    rule it produced two near-touching lines — a trailing-only rule reads
    just as clearly as a boxed footer without that artifact.)
    """
    parsed = [[s.strip() for s in line.split('|')] for line in lines]
    max_cols = max(len(segs) for segs in parsed)
    rows_html = []
    for idx, segments in enumerate(parsed):
        cells = []
        for i, seg in enumerate(segments):
            m = _META_FIELD_RE.match(seg)
            label, value = m.group(1).strip(), m.group(2).strip()
            # A leading bold tag like "**Footer:**" is sometimes just a section
            # label prefixed onto the first real field rather than a field of
            # its own (e.g. "**Footer:** Collection Executive: John") — when
            # the value itself parses as another Label: Value pair, prefer
            # that inner pair so "Collection Executive" gets its own clean
            # field instead of being nested as a value under "Footer".
            inner = _META_FIELD_RE.match(value)
            if inner and inner.group(1).strip():
                label, value = inner.group(1).strip(), inner.group(2).strip()
            attrs = ' style="width: 40%"' if i == 0 and len(segments) > 1 else ''
            # A row with fewer fields than the widest row in this group (e.g.
            # a lone "Ultrasound Findings: N/A" alongside two-field rows)
            # needs colspan on its last cell so its border — if this ends up
            # being the group's last row — stretches the full row width
            # instead of stopping at the first column.
            if i == len(segments) - 1 and len(segments) < max_cols:
                attrs += f' colspan="{max_cols - len(segments) + 1}"'
            cells.append(f'<td class="meta-cell"{attrs}><span class="meta-label">{label}:</span> {value}</td>')
        cls = ' class="meta-last"' if idx == len(parsed) - 1 else ''
        rows_html.append(f'<tr{cls}>{"".join(cells)}</tr>')
    return f'<table class="meta-table">{"".join(rows_html)}</table>'


def _is_section_heading(line: str) -> bool:
    s = line.strip()
    return s.startswith('**') and s.endswith('**') and s.count('**') == 2 and ':' not in s[2:-2] and '|' not in s


def _preprocess_pdf_text(text: str) -> tuple[str, dict[str, str]]:
    """
    Pulls patient-info/footer 'Label: Value' lines and standalone bold
    section titles out of the normal Markdown paragraph flow and pre-renders
    them (info lines as an aligned table row, section titles as a styled
    heading) so the PDF shows a real form layout instead of a single run-on
    line of text with literal pipe characters. Each becomes a placeholder
    token here, substituted back into the real HTML after markdown.markdown()
    runs, since injecting raw HTML mid-stream would otherwise fight the
    Markdown parser's own paragraph/emphasis handling.
    """
    placeholders: dict[str, str] = {}
    out_lines: list[str] = []
    meta_buffer: list[str] = []

    def flush_meta():
        if meta_buffer:
            key = f'KEPPLERPLACEHOLDER{len(placeholders)}X'
            placeholders[key] = _render_meta_group(list(meta_buffer))
            out_lines.append(key)
            meta_buffer.clear()

    for line in text.split('\n'):
        stripped = line.strip()
        if not stripped or stripped.startswith(('|', '#', '- ', '* ')):
            flush_meta()
            out_lines.append(line)
            continue

        if _is_section_heading(stripped):
            flush_meta()
            key = f'KEPPLERPLACEHOLDER{len(placeholders)}X'
            placeholders[key] = f'<h4 class="section-heading">{stripped.strip("*")}</h4>'
            out_lines.append(key)
            continue

        segments = [s.strip() for s in stripped.split('|')]
        if segments and all(_is_meta_field_segment(s) for s in segments):
            meta_buffer.append(stripped)
            continue

        flush_meta()
        out_lines.append(line)
    flush_meta()
    return '\n'.join(out_lines), placeholders


# Floor stops narrow columns collapsing below cell padding limits.
# Ceiling allows verbose columns (Impression, Findings) to take up to 60%.
MIN_COL_PCT = 8.0
MAX_COL_PCT = 60.0

# Preserved so the two tuned forms don't regress behind the generic path.
BLUEPRINT_OVERRIDES = {
    "LDSL": {
        "s.no": 8.0, "test description": 45.0,
        "sample type": 30.0, "lab name": 17.0,
    },
    "Healmax": {
        "s.no": 5.0, "patient name": 14.0, "age/sex": 8.0,
        "test code/name": 20.0, "sample type": 10.0, "barcode no": 10.0,
        "date/time": 12.0, "customer": 12.0, "referral doctor": 9.0,
    },
}

COMMON_HEADER_OVERRIDES = {
    "impression": 50.0,
    "findings": 50.0,
    "key findings": 50.0,
    "modality": 25.0,
    "study": 25.0,
    "investigation": 30.0,
    "date": 14.0,
    "time": 12.0,
    "s.no": 6.0,
    "sl.no": 6.0,
    "drug": 35.0,
    "medication": 35.0,
    "dose": 15.0,
    "route": 12.0,
    "freq": 15.0,
    "frequency": 15.0,
}


def _is_separator(line: str) -> bool:
    return bool(re.fullmatch(r'[\s|:\-]+', line.strip())) and '-' in line


def parse_markdown_table(md_text: str) -> list[list[str]]:
    """Return the widest markdown table in the text as a list of rows."""
    tables = parse_all_markdown_tables(md_text)
    if not tables:
        return []
    return max(tables, key=lambda t: max((len(r) for r in t), default=0))


def parse_all_markdown_tables(md_text: str) -> list[list[list[str]]]:
    """Return all markdown tables in the text as a list of table rows."""
    tables, current = [], []
    for line in md_text.split('\n'):
        stripped = line.strip()
        if stripped.startswith('|'):
            if _is_separator(stripped):
                continue
            current.append([c.strip() for c in stripped.strip('|').split('|')])
        elif current:
            tables.append(current)
            current = []
    if current:
        tables.append(current)
    return tables


def compute_column_widths(rows: list[list[str]]) -> list[float]:
    """
    Content-proportional widths, sqrt-dampened, clamped, renormalised to 100%.

    sqrt damping is the key choice: without it a 40-char column next to four
    2-char columns takes ~70% of the table and crushes the rest — exactly the
    FBS-vs-PLI/PD/U/INC failure. sqrt(40)/sqrt(2) is 4.5x, not 20x.
    """
    if not rows:
        return []
    ncols = max(len(r) for r in rows)
    if ncols == 0:
        return []

    weights = []
    for c in range(ncols):
        lengths = [len(r[c]) if c < len(r) else 0 for r in rows]
        header_len = lengths[0] if lengths else 0
        body_max = max(lengths[1:]) if len(lengths) > 1 else 0
        # Header still needs room even when the body is empty — an empty
        # column under a header called "ORDERED BY" must not go to zero.
        raw = max(header_len * 0.8, body_max, 1)
        weights.append(math.sqrt(raw))

    total = sum(weights)
    pcts = [w / total * 100.0 for w in weights]

    # Clamp, then redistribute the surplus across unclamped columns.
    for _ in range(10):
        clamped = [min(max(p, MIN_COL_PCT), MAX_COL_PCT) for p in pcts]
        drift = sum(clamped) - 100.0
        if abs(drift) < 0.01:
            pcts = clamped
            break
        free = [i for i, p in enumerate(clamped)
                if MIN_COL_PCT < p < MAX_COL_PCT]
        if not free:
            pcts = [100.0 / ncols] * ncols
            break
        share = drift / len(free)
        pcts = [p - share if i in free else p for i, p in enumerate(clamped)]
    else:
        pcts = [min(max(p, MIN_COL_PCT), MAX_COL_PCT) for p in pcts]

    # Round first, then absorb float residue into an unclamped column so MAX_COL_PCT is strictly respected.
    pcts = [round(p, 2) for p in pcts]
    residue = round(100.0 - sum(pcts), 2)
    if residue != 0:
        unclamped = [i for i, p in enumerate(pcts) if p < MAX_COL_PCT]
        target_idx = unclamped[0] if unclamped else pcts.index(max(pcts))
        pcts[target_idx] = round(pcts[target_idx] + residue, 2)
    return pcts


def _override_for(client_name: str, header: str) -> float | None:
    h_norm = header.strip().lower().replace('.', '').replace(' ', '')
    for key, table in BLUEPRINT_OVERRIDES.items():
        if key in client_name:
            for label, pct in table.items():
                if label.replace('.', '').replace(' ', '').replace('/', '') in h_norm.replace('/', ''):
                    return pct
    for label, pct in COMMON_HEADER_OVERRIDES.items():
        if label.replace('.', '').replace(' ', '').replace('/', '') in h_norm.replace('/', ''):
            return pct
    return None


def inject_column_widths(html_body: str, md_text: str = "",
                         client_name: str = "") -> str:
    """
    Inspects every <table> in html_body directly, calculates column widths for its <th> headers,
    and injects inline style="width: N%" into every <th> so every column is at least 8% and total is 100%.
    """
    def _table_repl(match: re.Match) -> str:
        table_html = match.group(0)

        th_matches = list(re.finditer(r'<th[^>]*>(.*?)</th>', table_html, flags=re.S | re.I))
        if not th_matches:
            return table_html

        ncols = len(th_matches)
        headers = [re.sub(r'<[^>]+>', '', m.group(1)).strip() for m in th_matches]

        final = []
        for h in headers:
            override = _override_for(client_name, h)
            if override is not None:
                final.append(max(MIN_COL_PCT, override))
            else:
                final.append(None)

        known_sum = sum(w for w in final if w is not None)
        unknown_count = sum(1 for w in final if w is None)

        if unknown_count > 0:
            rem = max(0.0, 100.0 - known_sum)
            even_w = max(MIN_COL_PCT, rem / unknown_count)
            final = [w if w is not None else even_w for w in final]

        total = sum(final)
        if total > 0:
            final = [round(w / total * 100.0, 2) for w in final]
        else:
            final = [round(100.0 / ncols, 2)] * ncols

        residue = round(100.0 - sum(final), 2)
        if residue != 0:
            final[-1] = round(final[-1] + residue, 2)

        th_idx = 0
        def _th_repl(m: re.Match) -> str:
            nonlocal th_idx
            if th_idx >= len(final):
                return m.group(0)
            w = final[th_idx]
            th_idx += 1
            inner = m.group(1)
            return f'<th style="width: {w}%">{inner}</th>'

        return re.sub(r'<th[^>]*>(.*?)</th>', _th_repl, table_html, flags=re.S | re.I)

    return re.sub(r'<table[^>]*>.*?</table>', _table_repl, html_body, flags=re.S | re.I)


def table_type_scale(ncols: int) -> dict:
    """Font size and padding scale dynamically with column count."""
    if ncols <= 5:
        return {"orientation": "portrait", "font": "10px",
                "td_pad": "7px 8px", "th_pad": "8px"}
    if ncols <= 8:
        return {"orientation": "landscape", "font": "10px",
                "td_pad": "7px 8px", "th_pad": "8px"}
    if ncols <= 11:
        return {"orientation": "landscape", "font": "8.5px",
                "td_pad": "4px 5px", "th_pad": "5px"}
    return {"orientation": "landscape", "font": "7.5px",
            "td_pad": "3px 4px", "th_pad": "4px"}


def generate_pro_pdf(md_text: str, client_name: str) -> bytes | str:
    try:
        safe = re.sub(r'<\|.*?\|>', '', md_text)
        safe = safe.replace('<', '&lt;').replace('>', '&gt;')
        safe = safe.replace('\u2022', '*').replace('\u2713', '[x]')
        
        # Step 2: Unconditionally normalize tables for all documents before column counting
        safe = _normalize_markdown_tables(safe)

        is_form = "LDSL" in client_name or "Healmax" in client_name
        
        if is_form:
            safe, meta_placeholders = _preprocess_pdf_text(safe)
            subtitle = "Test Requisition Form"
        else:
            meta_placeholders = {}
            subtitle = "Extraction Report"
            safe = re.sub(r'(?m)^(\s*)[-*]\s+', r'\1&#45; ', safe)
            
        ncols = _widest_table_column_count(safe)
        scale = table_type_scale(ncols)
        orientation = scale["orientation"]
        font_size = scale["font"]
        td_padding = scale["td_pad"]
        th_padding = scale["th_pad"]

        pdf_css = f"""
        @page {{ size: A4 {orientation}; margin: 1.8cm 1.5cm 1.8cm 1.5cm; }}
        body {{ font-family: Courier, monospace; font-size: 11px; color: #2a2a2a; line-height: 1.5; }}
        .raw-text {{ white-space: pre-wrap; }}
        .cover {{ padding: 10px 0 0; margin-bottom: 14px; }}
        .cover p.title {{ font-size: 22px; color: #123049; font-weight: bold; margin: 0; letter-spacing: 0.2px; font-family: Arial, sans-serif; }}
        .cover p.subtitle {{ font-size: 11px; color: #667080; margin: 3px 0 0 0; padding-bottom: 11px; border-bottom: 2px solid #d8dde2; font-family: Arial, sans-serif; }}
        h1 {{ font-size: 16px; color: #333; border-bottom: 1px solid #d8dde2; padding-bottom: 3px; margin-top: 18px; margin-bottom: 6px; font-family: Arial, sans-serif; }}
        h2 {{ font-size: 13px; color: #444; margin-top: 14px; margin-bottom: 6px; font-family: Arial, sans-serif; }}
        h3 {{ font-size: 11px; color: #555; margin-top: 10px; margin-bottom: 4px; font-family: Arial, sans-serif; }}
        table {{ width: 100%; table-layout: fixed; border-collapse: collapse; margin: 11px 0; font-size: {font_size}; font-family: Arial, sans-serif; }}
        thead {{ display: table-header-group; }}
        tr {{ page-break-inside: avoid; }}
        th {{ background: #eef2f7; color: #123049; padding: {th_padding}; text-align: left; font-weight: bold; border: 1px solid #dde2e7; word-wrap: break-word; overflow-wrap: break-word; }}
        td {{ border: 1px solid #dde2e7; padding: {td_padding}; vertical-align: top; word-wrap: break-word; }}
        tr:nth-child(even) td {{ background: #f8fafb; }}
        p {{ margin: 4px 0; }}
        ul {{ margin: 4px 0 8px 16px; padding: 0; }}
        li {{ margin: 2px 0; }}
        strong {{ color: #123049; }}
        .meta-table {{ width: 100%; table-layout: auto; border-collapse: collapse; margin: 6px 0; font-family: Arial, sans-serif; }}
        .meta-table td.meta-cell {{ border: none; padding: 4px 10px 4px 0; font-size: 10px; }}
        .meta-table tr.meta-last td.meta-cell {{ border-bottom: 1px solid #d8dde2; padding-bottom: 9px; }}
        .meta-label {{ font-weight: bold; color: #123049; }}
        .section-heading {{ font-size: 11.5px; font-style: italic; font-weight: bold; color: #123049; margin-top: 14px; margin-bottom: 5px; border-bottom: 1px solid #d8dde2; padding-bottom: 3px; font-family: Arial, sans-serif; }}
        .footer-note {{ font-size: 9px; color: #8a93a0; text-align: center; margin-top: 16px; border-top: 1px solid #e2e6ea; padding-top: 6px; font-family: Arial, sans-serif; }}
        """

        cover = f"""
        <div class="cover">
            <p class="title">{client_name}</p>
            <p class="subtitle">{subtitle}</p>
        </div>
        """

        footer = f"""
        <div class="footer-note">
            Generated on {time.strftime('%d %B %Y at %I:%M %p')} by Keppler AI. Verify all data against original records before use.
        </div>
        """
        
        safe = re.sub(r'([^\|\n][ \t]*)\n([ \t]*\|)', r'\1\n\n\2', safe)
        html_body = markdown.markdown(safe, extensions=['tables', 'nl2br'])

        for key, block_html in meta_placeholders.items():
            html_body = html_body.replace(key, block_html)
        block_alt = r'(?:<table class="meta-table">.*?</table>|<h4 class="section-heading">.*?</h4>)'
        html_body = re.sub(
            rf'<p>\s*({block_alt}(?:\s*<br\s*/?>\s*{block_alt})*)\s*</p>',
            lambda m: re.sub(r'<br\s*/?>', '', m.group(1)),
            html_body, flags=re.S,
        )

        # Step 1: Inject generic inline column widths for ALL tables
        html_body = inject_column_widths(html_body, safe, client_name)

        html = f"<html><head><meta charset='utf-8'/><style>{pdf_css}</style></head><body>{cover}{html_body}{footer}</body></html>"
        
        # Step 6: Try WeasyPrint first, fall back to xhtml2pdf/pisa
        try:
            from weasyprint import HTML, CSS
            weasy_css = pdf_css.replace("table-layout: fixed;", "table-layout: auto;")
            pdf_bytes = HTML(string=html).write_pdf(stylesheets=[CSS(string=weasy_css)])
            return pdf_bytes
        except Exception as weasy_err:
            logger.info(f"WeasyPrint rendering unavailable ({weasy_err}), using xhtml2pdf fallback")

        out = io.BytesIO()
        status = pisa.CreatePDF(io.StringIO(html), dest=out, encoding="utf-8")
        if status.err:
            raise RuntimeError(f"PDF generation failed: {status.err}")
        return out.getvalue()
    except Exception as e:
        raise RuntimeError(f"PDF Error: {e}")


def generate_docx(md_text: str, client_name: str) -> bytes | str:
    try:
        doc = Document()
        
        # Page margins
        for section in doc.sections:
            section.top_margin    = Inches(0.8)
            section.bottom_margin = Inches(0.8)
            section.left_margin   = Inches(1.0)
            section.right_margin  = Inches(1.0)
            
        # Cover heading
        title = doc.add_heading(f"Extraction Report: {client_name}", 0)
        title.alignment = WD_ALIGN_PARAGRAPH.LEFT
        title.runs[0].font.color.rgb = RGBColor(0x33, 0x33, 0x33)
        
        date_p = doc.add_paragraph(
            f"Generated: {time.strftime('%d %B %Y at %I:%M %p')}"
        )
        date_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        date_p.runs[0].font.size = Pt(9)
        date_p.runs[0].font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        doc.add_paragraph()
        
        clean = re.sub(r'<\|.*?\|>', '', md_text)
        in_table = False
        table    = None
        
        for line in clean.split('\n'):
            line = line.strip()
            if not line:
                if not in_table:
                    doc.add_paragraph()
                continue
            
            if line.startswith('## '):
                in_table = False
                h = doc.add_heading(line[3:], level=1)
                h.runs[0].font.color.rgb = RGBColor(0x1a, 0x52, 0x76)
            elif line.startswith('### '):
                in_table = False
                h = doc.add_heading(line[4:], level=2)
                h.runs[0].font.color.rgb = RGBColor(0x1a, 0x52, 0x76)
            elif line.startswith('- ') or line.startswith('* '):
                in_table = False
                p = doc.add_paragraph(style='List Bullet')
                text = line[2:].strip()
                parts = re.split(r'\*\*(.+?)\*\*', text)
                for i, part in enumerate(parts):
                    run = p.add_run(part)
                    if i % 2 == 1: run.bold = True
            elif line.startswith('|'):
                if _is_table_separator(line):
                    continue
                cells = [c.strip() for c in line.strip('|').split('|')]
                cells = [c for c in cells if c or len(cells) > 1] # allow empty cells if part of row
                if not cells:
                    continue
                if not in_table:
                    table = doc.add_table(rows=1, cols=len(cells))
                    table.style = 'Table Grid'
                    hdr = table.rows[0]
                    for i, c in enumerate(cells):
                        cell = hdr.cells[i]
                        cell.text = c
                        run = cell.paragraphs[0].runs[0] if cell.paragraphs[0].runs else cell.paragraphs[0].add_run(c)
                        run.bold = True
                        cell._tc.get_or_add_tcPr()
                    in_table = True
                else:
                    row = table.add_row()
                    for i, c in enumerate(cells):
                        if i < len(row.cells):
                            row.cells[i].text = c
            elif line.startswith('---'):
                in_table = False
            else:
                in_table = False
                p = doc.add_paragraph()
                parts = re.split(r'\*\*(.+?)\*\*', line)
                for i, part in enumerate(parts):
                    run = p.add_run(part)
                    if i % 2 == 1: run.bold = True
                    run.font.size = Pt(10)
        
        doc.add_paragraph()
        note = doc.add_paragraph(
            "⚠ This extraction was auto-generated by Keppler AI from scanned case file documents. "
            "Verify all clinical data against original records before use."
        )
        note.runs[0].font.size = Pt(8)
        note.runs[0].font.color.rgb = RGBColor(0x88, 0x88, 0x88)
        
        out = io.BytesIO()
        doc.save(out)
        return out.getvalue()
    except Exception as e:
        return f"Word Error: {e}"

def generate_excel(md_text: str) -> bytes | None:
    if not md_text.strip():
        return None
        
    try:
        out = io.BytesIO()
        with pd.ExcelWriter(out, engine='openpyxl') as writer:
            # --- 1. CONSOLIDATED REPORT (Primary Sheet) ---
            report_data = []
            lines = md_text.split('\n')
            
            # Extract key-value pairs (like **Key:** Value)
            metadata = []
            other_text = []
            tables = []
            current_table = []
            
            for line in lines:
                raw_line = line.strip()
                if not raw_line: continue
                
                # Table detection — require real Markdown table syntax (leading '|'),
                # not just 2+ pipes. The LDSL/Healmax blueprints also use '|' as an
                # informal separator in plain text lines (e.g. "**Age:** 34 | **Sex:** F"),
                # which a "count('|') >= 2" check would misclassify as a table row.
                if raw_line.startswith('|'):
                    if _is_table_separator(raw_line):
                        continue
                    cells = [c.strip() for c in raw_line.split('|') if c.strip()]
                    if cells:
                        current_table.append(cells)
                    continue
                else:
                    if current_table:
                        if len(current_table) > 1:
                            tables.append(current_table)
                        current_table = []

                # Metadata detection (Bold keys) — split on '|' first so a single line
                # packing several "**Key:** Value" fields (as LDSL/Healmax headers do)
                # becomes one metadata row per field instead of one merged blob.
                if '**' in raw_line and ':' in raw_line:
                    for segment in raw_line.split('|'):
                        segment = segment.strip()
                        if not segment:
                            continue
                        parts = re.split(r'\*\*|\*', segment)
                        cleaned = " ".join([p.strip() for p in parts if p.strip()])
                        if cleaned:
                            metadata.append([cleaned])
                else:
                    other_text.append([raw_line])
            
            if current_table and len(current_table) > 1:
                tables.append(current_table)

            # Build the master list for the first sheet
            # [A] Patient / Form Metadata
            if metadata:
                report_data.append(["--- FORM DETAILS ---"])
                report_data.extend(metadata)
                report_data.append([""]) # Spacer
            
            # [B] Tables
            if tables:
                for i, table in enumerate(tables):
                    report_data.append([f"--- TABLE {i+1} ---"])
                    report_data.extend(table)
                    report_data.append([""]) # Spacer
            
            # [C] Other Text / Footer
            if other_text:
                report_data.append(["--- ADDITIONAL INFORMATION ---"])
                report_data.extend(other_text)
            
            # Create the Consolidated Sheet
            # Find max width to avoid dataframe errors
            max_cols = max([len(r) for r in report_data]) if report_data else 1
            df_cons = pd.DataFrame(report_data, columns=[f"Col {i+1}" for i in range(max_cols)])
            df_cons.to_excel(writer, index=False, header=False, sheet_name='Consolidated Report')
            
            # --- 2. CLEAN TABLE SHEETS (Individual sheets for analysis) ---
            if tables:
                for i, table in enumerate(tables):
                    header = table[0]
                    rows = table[1:]
                    cleaned_rows = [r[:len(header)] if len(r) > len(header) else r + ['']*(len(header)-len(r)) for r in rows]
                    df_table = pd.DataFrame(cleaned_rows, columns=header)
                    sheet_name = f'Table_{i+1}'
                    df_table.to_excel(writer, index=False, sheet_name=sheet_name)
            
            # --- 3. RAW EXTRACTION SHEET ---
            paragraphs = [p.strip() for p in md_text.split('\n') if p.strip()]
            df_full = pd.DataFrame({"Raw Content": paragraphs})
            df_full.to_excel(writer, index=False, sheet_name='Raw Text')
            
        return out.getvalue()
    except Exception as e:
        print(f"Excel Error: {e}")
        # Final fallback
        try:
            out = io.BytesIO()
            pd.DataFrame({"Extraction": [md_text]}).to_excel(out, index=False)
            return out.getvalue()
        except:
            return None

def generate_json(md_text: str) -> str:
    if not md_text.strip(): return "{}"
    
    data = {"metadata": {}, "tables": [], "text": []}
    lines = md_text.split('\n')
    current_table = []
    
    for line in lines:
        raw = line.strip()
        if not raw: continue
        
        if raw.startswith('|'):
            if _is_table_separator(raw): continue
            cells = [c.strip() for c in raw.split('|') if c.strip()]
            if cells: current_table.append(cells)
            continue
        else:
            if current_table and len(current_table) > 1:
                header = current_table[0]
                rows = current_table[1:]
                table_data = []
                for row in rows:
                    row_dict = {}
                    for i, h in enumerate(header):
                        row_dict[h] = row[i] if i < len(row) else ""
                    table_data.append(row_dict)
                data["tables"].append(table_data)
                current_table = []

        if '**' in raw and ':' in raw:
            for segment in raw.split('|'):
                segment = segment.strip()
                if not segment:
                    continue
                parts = segment.split('**')
                if len(parts) >= 3:
                    key = parts[1].replace(':', '').strip()
                    val = "".join(parts[2:]).strip()
                    if key:
                        data["metadata"][key] = val
                elif segment:
                    data["text"].append(segment)
        else:
            data["text"].append(raw)
            
    if current_table and len(current_table) > 1:
        header = current_table[0]
        rows = current_table[1:]
        table_data = []
        for row in rows:
            row_dict = {}
            for i, h in enumerate(header):
                row_dict[h] = row[i] if i < len(row) else ""
            table_data.append(row_dict)
        data["tables"].append(table_data)
        
    return json.dumps(data, indent=2)

# ---------------------------------------------------------------------------
# 6.  BLUEPRINTS
# ---------------------------------------------------------------------------
BLUEPRINTS = {
    "Universal OCR (Any Text)": {
        "identity": "You are a state-of-the-art Universal OCR Vision Model, operating at the same capability level as Google Vision and OpenAI GPT-4V. You possess vast world knowledge across all domains (medical, legal, financial, etc.).",
        "structure": "Organize the extracted text into a clean, professional Markdown format. Use Markdown Headers (###) for sections, **bold text** for keys/labels, bullet points (-) for lists, and Markdown Tables (|---|) for any tabular or grid-like data.",
        "instructions": "MANDATORY: Extract ANY kind of text, across different font styles, variations, and especially messy cursive handwriting with 100% zero-shot accuracy. If the document is handwritten (e.g. a medical prescription), transcribe literally. Look for standard abbreviations like 'Tab', 'Cap', 'Syr'. Do not output your internal reasoning.",
        "rules": "2. EXHAUSTIVE TRANSCRIPTION: Do not stop generating until the very last word of the image is transcribed.\n3. LITERAL TRANSCRIPTION: Transcribe the glyphs you actually see. Do NOT substitute a plausible drug name for an implausible one. Do NOT expand abbreviations (write ATORVAS, not atorvastatin). Do NOT normalise brand to generic. If a character is genuinely ambiguous, transcribe your best reading wrapped in ?question marks? — never silently pick one.\n4. NO PREAMBLE: Start directly with the extracted data. Do not output the document domain or any greetings.\n5. FORMATTING: Use Markdown to give the text a professional structure."
    },
    "Drug Chart / MAR": {
        "identity": (
            "You are a medication-chart transcription engine reading an inpatient "
            "Drug Chart / Medication Administration Record (MAR). You transcribe "
            "literally and flag uncertainty. You do not interpret, expand, or "
            "clinically reason about what is written."
        ),
        "structure": """\
**Patient Name:** [name] | **Age/Sex:** [as printed] | **Admission No:** [no]
**UHID No:** [no] | **Consultant:** [name] | **Admission Date:** [DD MMM YYYY]
**Drug Allergy:** [as written, e.g. Nil] | **Physician:** [name in the "To be filled by the Physician" box]

| Date | Time | Name of the Drug | Dose | Route | Freq. | Doctor Name & Sign | Time given (T) / Initials of Nurse (I) | Stop Orders (Doctor Name & Sign / With Time) |
|------|------|------------------|------|-------|-------|--------------------|----------------------------------------|---------------------------------------------|

**Footer:** [form number / revision / issue date as printed]
**Handwritten notes below table:** [any order written below the printed grid, transcribed literally]""",
        "instructions": (
            "This form has NINE columns under THREE printed column groups. Group 1 "
            "'To be filled by the Physician (IN CAPITAL LETTERS)' contains Date, Time, "
            "Name of the Drug, Dose, Route, Freq., and Doctor Name & Sign. Group 2 "
            "'To be filled by Nurse' contains Time given (T) / Initials of the Nurse (I). "
            "Group 3 'Stop Orders' contains Doctor Name & Sign and With Time. "
            "These groups are SEPARATE. A time written in the nurse's column is an "
            "administration time, NOT a stop order. Never move a value across a column "
            "boundary."
        ),
        "rules": (
            "2. LITERAL TRANSCRIPTION: Transcribe the glyphs you actually see. Do NOT "
            "substitute a plausible drug name for an implausible one. Do NOT expand "
            "abbreviations — write ATORVAS, not atorvastatin; write INJ, not Injection. "
            "Do NOT normalise brand names to generic names. If a character is genuinely "
            "ambiguous, transcribe your best reading wrapped in ?question marks? — never "
            "silently pick one. On a medication chart a confident wrong name is more "
            "dangerous than a flagged unreadable one.\n"

            "3. FREQUENCY IS A CLOSED SET. The Freq. column may only contain one of: "
            "ALT, BD, HS, OD, Q12H, Q4H, Q6H, Q8H, QDS, QID, SOS, STAT, TDS, TID, WEEKLY. Read each row's Freq. cell "
            "INDEPENDENTLY of every other row. Never copy the value from the row above. "
            "Never assume a default such as BD. If unreadable write ?. If empty leave empty.\n"

            "4. UNITS ARE TRANSCRIBED, NOT CONVERTED. Write exactly the unit on the page: "
            "mg, mcg, g, ml, IU, amp, tab. mcg and mg are DIFFERENT — look carefully at "
            "which is written and never substitute one for the other. Never convert "
            "between units. Never add a unit that is not written.\n"

            "5. EMPTY CELLS STAY EMPTY. If a cell has nothing in it, output an empty "
            "cell. Do not write N/A, None, or '-'. Do not carry a value down from the "
            "row above — a date written once at the top of a merged cell applies to that "
            "cell only; mark inherited dates as (inherited).\n"

            "6. COLUMN INTEGRITY. Each of the nine columns is read independently. A "
            "value belongs to the column its glyphs physically sit inside. If a "
            "handwritten stroke crosses a printed rule line, assign it to the column "
            "containing most of the stroke and flag it with ?.\n"

            "7. PREFIXES. INJ, TAB, CAP, SYP are dosage-form prefixes and belong in the "
            "Name of the Drug column. Transcribe them as written. 'INJ' in cursive is "
            "commonly misread as 'BNT' or 'DNT' — if you see what looks like BNT before "
            "a drug name, it is almost certainly INJ, but transcribe what is legible and "
            "flag it rather than guessing the rest of the name to match.\n"

            "8. NO PREAMBLE. Start directly with the Patient Name line. No commentary.\n"

            "9. COMPLETENESS. Output every row of the grid that has any content, the "
            "header block above the grid, and any handwritten order below the grid. Do "
            "not stop early."
        ),
    },
    "LDSL Diagnostics": {
        "identity": "You are a high-performance medical OCR engine (Qwen 2.5 VL optimized), specialized in reading LDSL Diagnostics test requisition forms.",
        "structure": """\
**Patient Name:** [Full Name as written, preserve letter spacing if boxed] | **Age:** [Age] | **Sex:** [M/F]
**Referred Doctor:** [Doctor Name] | **Collection Date:** [DD-MM-YYYY] | **Collection Time:** [HH:MM AM/PM]

| S.No | Test Description | Sample Type (Please Tick) | Lab Name / C.C. Code |
|------|-------------------|----------------------------|------------------------|
| 1    | [Test Name]       | [Ticked/circled sample type, exactly as marked] | [Code if printed, else blank] |

**History for Quadruple/Triple/Double Marker Tests:**
Date of Birth: [DOB] | Weight: [Weight]
LMP: [LMP] | Gestation Age: [Gestation Age]
Diabetes Status: [Yes/No/N/A] | Gestation: [Single/Twin]
Ultrasound Findings: [Details]

**Footer:** Collection Executive: [Name] | Date: [DD-MM-YYYY] | Checked by: [Name]""",
        "instructions": "MANDATORY: Transcribe every handwritten entry exactly. Use the 4-column Markdown table for all 10 test rows — do NOT merge the 'Sample Type' and 'Lab Name / C.C. Code' columns, they are always separate even if visually close together. Do not omit the Patient Name, History, or Footer sections.",
        "rules": "2. SPATIAL AWARENESS: Maintain the exact layout and hierarchy seen in the document — S.No, Test Description, Sample Type, and Lab Name/C.C. Code are four distinct columns, never combine them into one cell.\n3. HANDWRITING & TICKS: Transcribe every handwritten scribble or mark. If a sample type box is ticked/circled, write only that ticked option; if none is marked, leave the cell blank rather than guessing.\n4. NUMERICAL PRECISION: Do not round or alter any numbers, dates, or measurements. If a date or time is partially cut off or ambiguous, transcribe exactly what is visible rather than inferring the missing part.\n5. NO PREAMBLE: Start directly with the extracted data. No greetings or meta-commentary.\n6. TABLE INTEGRITY: Ensure all 10 test rows are present and every column is populated based on the visual rows, in reading order top to bottom.\n7. MISSING DATA: If any field is missing from the document, leave it completely blank. Do NOT write 'Not documented', 'N/A', or 'None' unless that literal text is printed on the form."
    },
    "Healmax Diagnostics": {
        "identity": "You are a high-performance medical OCR engine (Qwen 2.5 VL optimized).",
        "structure": """\
**Franchisee Code:** ___ | **Date:** ___

| S.No | Patient Name | Age/Sex | Test Code/Name | Sample Type | Barcode No | Date/Time | Customer | Referral Doctor |
|------|-------------|---------|----------------|-------------|------------|-----------|----------|-----------------|""",
        "instructions": "Fill all 9 table columns. Do NOT merge or skip any column.",
        "rules": "2. SPATIAL AWARENESS: Maintain the exact layout and hierarchy seen in the document.\n3. HANDWRITING: Transcribe every handwritten scribble or mark. If a checkmark is present in a box, represent it as [x].\n4. NUMERICAL PRECISION: Do not round or alter any numbers, dates, or measurements.\n5. NO PREAMBLE: Start directly with the extracted data. No greetings or meta-commentary.\n6. TABLE INTEGRITY: Ensure every column of the table is populated correctly based on the visual rows.\n7. MISSING DATA: If any field is missing from the document, leave it completely blank. Do NOT write 'Not documented', 'N/A', or 'None'."
    },
    "Handwritten Medical Prescription": {
        "identity": "You are an expert Clinical Pharmacist AI. You specialize in deciphering highly illegible, cursive doctor handwriting.",
        "structure": "Organize the extracted text into a clean Markdown format. Use bullet points for medications, including dosages and frequencies if present.",
        "instructions": "MANDATORY: Transcribe literally. Look for standard abbreviations like 'Tab' (Tablet), 'Cap' (Capsule), 'Syr' (Syrup), 'Inj' (Injection).",
        "rules": "2. LITERAL TRANSCRIPTION: Transcribe the glyphs you actually see. Do NOT substitute a plausible drug name for an implausible one. Do NOT expand abbreviations. Do NOT normalise brand to generic. If a character is genuinely ambiguous, transcribe your best reading wrapped in ?question marks? — never silently pick one.\n3. NO PREAMBLE: Start directly with the extracted data.\n4. FIDELITY: Do not hallucinate generic names if a brand name is written."
    },
    "Nursing Handover / ISBAR Checklist": {
        "identity": "You are a clinical transcription engine reading a nursing shift-handover checklist grid (ISBAR / handing-over-taking-over chart). The form is a ruled grid: printed check items run DOWN the left column, and each column to the right is one handover entry (one shift).",
        "structure": "A SINGLE markdown table. First column = the printed check item, copied verbatim including its '(Y/N)' or '(Y/N/NA)' suffix. Then one column per printed entry column on the form, named 'Entry 1', 'Entry 2', ... Emit every printed check-item row, in printed order, even when the whole row is blank.",
        "instructions": "MANDATORY: Output ONE table. Never split the grid into several tables. Never promote a check item into a column header. Count the ruled vertical lines to determine the number of entry columns and emit that many columns for every row — a blank entry column is data and must be emitted as an empty cell.",
        "rules": (
            "2. CLOSED VALUE SET: each entry cell is one of Y, N, NA, or empty. If a handwritten mark is ambiguous between Y and N, output [?] for that cell. Never guess between Y and N — they are clinically opposite.\n"
            "3. Y vs N: a looped/curved single stroke resembling a lowercase \"y\" with a descender is Y. An angular two-stroke or \"n\"-like glyph is N. If the stroke has neither a clear descender nor a clear angular joint, output [?].\n"
            "4. A blank cell means \"not recorded\". Never carry a value across from the column to the left or the row above.\n"
            "5. TIME AND SIGNATURE ROWS: \"Completion Time\" cells are times written as given (e.g. \"8:30 am\"). \"Handing over Nurse Sign & ID No.\" and \"Taking over Nurse Sign & ID No.\" are SIGNATURE cells — output the hand-PRINTED name if one is legible, otherwise leave the cell EMPTY. Never write [UNCLEAR] or a description of a squiggle in a signature cell.\n"
            "6. NEVER pad the table with extra blank rows to make it look complete. The table ends at the last printed check item."
        )
    },
    "Procedure Safety Checklist": {
        "identity": "You are a clinical transcription engine extracting header metadata and signature details from a Procedure Safety Checklist form.",
        "structure": "Output ONLY the header fields (Date, IP No, Patient name, Age/gender, Doctor Name), time fields, and printed signature names in clean Markdown key-value lines.",
        "instructions": (
            "Do NOT transcribe the checklist items. Do NOT list the checkbox options. "
            "The checklist is extracted separately by a dedicated detector. "
            "Output ONLY the header fields (Date, IP No, Patient name, Age/gender, Doctor Name), time fields, and hand-PRINTED signature values."
        ),
        "rules": "1. NO PREAMBLE: Start directly with the extracted metadata.\n2. NO CHECKLIST ITEMS: Do NOT generate or list any checklist item text or checkboxes."
    }
}

# ---------------------------------------------------------------------------
# 7.  PROMPT BUILDER
# ---------------------------------------------------------------------------
def build_prompt(client: str, page_info: str = "") -> str:
    client_upper = (client or "").upper()
    if any(kw in client_upper for kw in (
        "SAFETY CHECKLIST", "OPTHAL/06"
    )):
        bp = BLUEPRINTS["Procedure Safety Checklist"]
    elif any(kw in client_upper for kw in (
        "ISBAR", "HANDOVER", "HANDING OVER", "TAKING OVER",
        "SHIFT HANDOVER", "NURSE HANDOVER", "SAFETY CHECK"
    )):
        bp = BLUEPRINTS["Nursing Handover / ISBAR Checklist"]
    else:
        bp = BLUEPRINTS.get(client, BLUEPRINTS["Universal OCR (Any Text)"])



    page_note = f" ({page_info})" if page_info else ""
    return f"""{bp['identity']}{page_note}

Your objective: Extract EVERY piece of text from the image with 100% fidelity.

Output Format:
{bp['structure']}

Rules for 100% Accuracy:
1. MANDATORY: {bp['instructions']}
{bp['rules']}"""


# ---------------------------------------------------------------------------
# 8.  CORE PIPELINE (Streamlit-free — used by both render_ocr_app() and the FastAPI worker)
# ---------------------------------------------------------------------------
def detect_vertical_rule_lines(raw_img: Image.Image) -> list[int]:
    """
    Detects vertical rule lines in a printed table image using OpenCV morphological filtering + HoughLinesP.
    Returns sorted X-coordinates of vertical lines.
    """
    try:
        import cv2
        import numpy as np
        img_cv = cv2.cvtColor(np.array(raw_img), cv2.COLOR_RGB2BGR)
        gray = cv2.cvtColor(img_cv, cv2.COLOR_BGR2GRAY)
        h, w = gray.shape
        kernel_length = max(10, h // 30)
        vert_kernel = cv2.getStructuringElement(cv2.MORPH_RECT, (1, kernel_length))
        vert_img = cv2.erode(gray, vert_kernel, iterations=2)
        vert_img = cv2.dilate(vert_img, vert_kernel, iterations=2)
        lines = cv2.HoughLinesP(vert_img, 1, np.pi/180, threshold=100, minLineLength=kernel_length, maxLineGap=10)
        
        if lines is None:
            return []
        
        x_coords = []
        for line in lines:
            line_val = line[0] if getattr(line, 'ndim', 1) > 1 else line
            if len(line_val) >= 4:
                x1, y1, x2, y2 = line_val[:4]
                if abs(x1 - x2) <= 5:
                    x_coords.append(int((x1 + x2) // 2))
        
        if not x_coords:
            return []
            
        x_coords.sort()
        unique_x = []
        for x in x_coords:
            if not unique_x or abs(x - unique_x[-1]) > 15:
                unique_x.append(x)
                
        return unique_x
    except Exception as e:
        logger.warning(f"Vertical line detection failed: {e}")
        return []


def process_single_page(
    raw_img: Image.Image,
    label: str = "Page 1",
    idx: int = 0,
    total_pages: int = 1,
    client: str = "Universal OCR (Any Text)",
    progress_cb=None,
) -> dict:
    """
    Process exactly one page image and return its text + predictions.
    Used by workers/celery_app.py for per-page Celery tasks.

    Returns {"label": str, "text": str, "predictions": list[dict]}.
    """
    if progress_cb: progress_cb(0.02)

    raw_img = apply_orientation(raw_img)
    q = assess_input_quality(raw_img)

    if progress_cb: progress_cb(0.05)


    page_info = f"{label} of {total_pages}" if total_pages > 1 else ""
    prompt = build_prompt(client, page_info)

    if progress_cb: progress_cb(0.15)
    
    async_ocr = AsyncRegionOCR(max_concurrent=5, model_name="qwen2.5-vl-7b")
    w, h = raw_img.size

    from modules.form_templates import match_template, build_fill_prompt, render_markdown
    from modules.grid_geometry import detect_grid_columns

    # Cheap probe text extraction (top 15% and bottom 8%)
    top_strip = raw_img.crop((0, 0, w, max(1, int(h * 0.15))))
    bot_strip = raw_img.crop((0, max(0, int(h * 0.92)), w, h))

    probe_top, _, _, _ = asyncio.run(async_ocr._call_model_with_retry_async(top_strip, "Read all text in this top header strip."))
    probe_bot, _, _, _ = asyncio.run(async_ocr._call_model_with_retry_async(bot_strip, "Read all text in this bottom footer strip."))
    probe_text = f"{probe_top}\n{probe_bot}"

    template = match_template(probe_text) or match_template(client)
    template_used = False

    is_safety_checklist = any(kw in (client + " " + probe_text).upper() for kw in (
        "SAFETY CHECKLIST", "OPTHAL/06", "SIGN IN", "TIME OUT", "SIGN OUT"
    ))

    if is_safety_checklist:
        from modules.checkbox_detector import detect_checkboxes, render_checked_only
        cbs = detect_checkboxes(raw_img)
        checked_md = render_checked_only(cbs)
        header_text, strategy_used, predictions, ocr_confidence = asyncio.run(
            async_ocr._call_model_with_retry_async(raw_img, prompt)
        )
        extracted_text = f"{header_text}\n\n{checked_md}"
        template_used = True
    elif template is not None:
        logger.info(f"Form template matched: {template.template_id}")

        cols = detect_grid_columns(raw_img)
        n_cols = len(cols) if len(cols) >= 3 else 10
        fill_prompt = build_fill_prompt(template, n_cols)

        raw_resp, strategy_used, predictions, ocr_confidence = asyncio.run(
            async_ocr._call_model_with_retry_async(raw_img, fill_prompt)
        )
        strategy_used = f"template:{template.template_id}"

        clean_resp = raw_resp.strip()
        if "```json" in clean_resp:
            clean_resp = clean_resp.split("```json")[1].split("```")[0].strip()
        elif clean_resp.startswith("```"):
            clean_resp = clean_resp.split("```")[1].strip()

        cells_dict = {}
        if clean_resp.startswith("{") and clean_resp.endswith("}"):
            try:
                parsed = json.loads(clean_resp)
                if isinstance(parsed, dict):
                    cells_dict = parsed.get("cells", {})
            except Exception as e:
                logger.warning(f"Failed to parse JSON cells from template fill: {e}")

        if cells_dict:
            rendered_md, tmpl_needs_review = render_markdown(template, cells_dict, n_cols)
            rendered_md, violations = enforce_template_labels(rendered_md, template)
            if violations:
                logger.error(f"FABRICATED LABELS REJECTED: {violations}")
                tmpl_needs_review = True
                rendered_md = ("> **[FABRICATED CONTENT REMOVED — "
                              f"{len(violations)} invented row(s) dropped. VERIFY AGAINST ORIGINAL]**\n\n") + rendered_md
            extracted_text = rendered_md
            template_used = True


    if not template_used:
        x_rules = detect_vertical_rule_lines(raw_img)
        if client == "Drug Chart / MAR" and len(x_rules) >= 6:
            # --- PATH C: Column-Band Grounded Mode for Drug Chart ---
            if progress_cb: progress_cb(0.30)
            logger.info(f"Drug Chart vertical line detection found {len(x_rules)} rules — running column-band OCR")

            drug_col_x1 = x_rules[min(2, len(x_rules)-1)]
            drug_col_x2 = x_rules[min(3, len(x_rules)-1)] if len(x_rules) > 3 else w

            drug_band_img = raw_img.crop((max(0, drug_col_x1 - 5), 0, min(w, drug_col_x2 + 5), h))
            drug_band_upscaled = drug_band_img.resize((drug_band_img.width * 2, drug_band_img.height * 2), Image.Resampling.LANCZOS)

            extracted_text, strategy_used, predictions, ocr_confidence = asyncio.run(
                async_ocr._call_model_with_retry_async(raw_img, prompt)
            )
            if progress_cb: progress_cb(0.80)
        else:
            # --- PATH B: Single-Shot Mode (Handwriting Context) ---
            if progress_cb: progress_cb(0.40)

            extracted_text, strategy_used, predictions, ocr_confidence = asyncio.run(
                async_ocr._call_model_with_retry_async(raw_img, prompt)
            )

    if progress_cb: progress_cb(0.85)


    if extracted_text:
        extracted_text = prune_empty_table_rows(extracted_text)
        extracted_text = deduplicate_markdown_tables(extracted_text)

    # Run Chart Validation and hard confidence gating (Step 6 & Step 7)

    needs_review = False
    validation_flags = []
    if extracted_text:
        rows = []
        for line in extracted_text.splitlines():
            if "|" in line and not line.strip().startswith("|---") and not line.strip().startswith("| ---"):
                parts = [p.strip() for p in line.split("|")[1:-1]]
                if len(parts) >= 6 and parts[0] != "Date":
                    rows.append({
                        "Date": parts[0],
                        "Time": parts[1] if len(parts) > 1 else "",
                        "drug_name": parts[2] if len(parts) > 2 else "",
                        "dose": parts[3] if len(parts) > 3 else "",
                        "route": parts[4] if len(parts) > 4 else "",
                        "freq": parts[5] if len(parts) > 5 else "",
                        "doctor": parts[6] if len(parts) > 6 else "",
                        "nurse": parts[7] if len(parts) > 7 else "",
                        "stop_orders": parts[8] if len(parts) > 8 else "",
                    })
        if rows:
            from modules.chart_validator import validate_chart_rows
            validation_flags = validate_chart_rows(rows)
            for flag in validation_flags:
                if flag.get("severity") == "critical":
                    needs_review = True

    if ocr_confidence < 0.85:
        needs_review = True

    page_preds = []
    if extracted_text:
        full_page_bbox = [0, 0, w, h]
        for pred in predictions:
            sem_conf = float(pred.get("Confidence", 0.0))
            final_conf = ConfidenceEngine.calculate_final_confidence(1.0, sem_conf)
            pred["Confidence"] = f"{final_conf:.2f}"
            pred["page"] = idx + 1
            pred["bbox"] = full_page_bbox
            pred["ocr_confidence"] = ocr_confidence
            pred["ocr_model_used"] = "qwen2.5-vl-7b"
            pred["region_id"] = "full-page"
            pred["region_type"] = "Full Page"
            pred["needs_review"] = needs_review
            pred["validation_flags"] = validation_flags
            page_preds.append(pred)
            
        page_text = extracted_text
        client_upper = (client or "").upper()
        if any(kw in client_upper for kw in (
            "ISBAR", "HANDOVER", "HANDING OVER", "TAKING OVER",
            "SHIFT HANDOVER", "NURSE HANDOVER", "SAFETY CHECK"
        )):
            from modules.required_fields import audit_grid_row_coverage, ISBAR_EXPECTED_LABELS
            missing = audit_grid_row_coverage(page_text, ISBAR_EXPECTED_LABELS)
            if missing:
                needs_review = True
                n_missing = len(missing)
                page_text = f"> **[MISSING ROWS — {n_missing} printed check items not transcribed]**\n\n{page_text}"
    else:
        page_text = f"*[{label}: Extraction failed — content too short or empty]*"


    if progress_cb: progress_cb(0.95)

    return {"label": label, "text": page_text, "predictions": page_preds}


def run_ocr_pipeline(pages: list[Image.Image], client: str, progress_cb=None) -> dict:
    """
    Sequential wrapper around process_single_page for direct/synchronous callers
    (tests, small documents processed in-process). The distributed path
    (workers/celery_app.py) calls process_single_page directly, one page per
    Celery task, instead of using this loop.

    progress_cb, if given, is called with a float in [0, 1] as pages complete.

    Returns {"pages": [(label, text), ...], "combined": str, "predictions": list[dict], "elapsed": float}.
    """
    start_time = time.time()
    total_pages = len(pages)
    page_labels = [f"Page {i+1}" for i in range(total_pages)]

    all_pages_text = []
    all_preds = []

    for idx, (raw_img, label) in enumerate(zip(pages, page_labels)):
        base_pct = idx / total_pages
        page_slice = 1.0 / total_pages

        def page_progress_cb(p, base_pct=base_pct, page_slice=page_slice):
            if progress_cb: progress_cb(base_pct + p * page_slice)

        result = process_single_page(raw_img, label, idx, total_pages, client, progress_cb=page_progress_cb)
        all_pages_text.append((result["label"], result["text"]))
        all_preds.extend(result["predictions"])

    if progress_cb:
        progress_cb(1.0)

    elapsed = time.time() - start_time

    if total_pages == 1:
        combined = all_pages_text[0][1]
    else:
        parts = [f"---\n### {lbl}\n\n{txt}" for lbl, txt in all_pages_text]
        combined = "\n\n".join(parts)

    return {
        "pages": all_pages_text,
        "combined": combined,
        "predictions": all_preds,
        "elapsed": elapsed,
    }

